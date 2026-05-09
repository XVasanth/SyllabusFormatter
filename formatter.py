"""
formatter.py
────────────
Core Word document formatting engine.
Accepts a settings dict from template_reader.py and applies all rules.

Features:
  - Page size (A4 / Letter / Legal)
  - Page border with independent cm size per edge
  - Per-level fonts: unit heading / topic heading / body /
                     table header / table body / references
  - Run-level classification via Groq (topic headings inside paragraphs)
  - Reference part formatting (author italic, title bold, etc.) via Groq
  - Lecture hours: bold, italic, right-aligned tab stop
  - Table cell borders
  - Groq-cleaned reference text (edition number normalisation)
"""

import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE SIZE
# ══════════════════════════════════════════════════════════════════════════════

_PAGE_SIZES = {
    "a4":     (Inches(8.27),  Inches(11.69)),
    "letter": (Inches(8.5),   Inches(11.0)),
    "legal":  (Inches(8.5),   Inches(14.0)),
}

def apply_page_size(doc, size_name: str):
    key = size_name.lower()
    if key not in _PAGE_SIZES:
        return
    w, h = _PAGE_SIZES[key]
    for section in doc.sections:
        section.page_width  = w
        section.page_height = h


# ══════════════════════════════════════════════════════════════════════════════
#  PAGE BORDER
# ══════════════════════════════════════════════════════════════════════════════

_CM_TO_TWENTIETHS = 567

def apply_page_border(doc, style: str, color: str, sizes: dict):
    if style == "none":
        return
    color = color.lstrip("#").upper() or "000000"
    val_map = {"single":"single","double":"double","shadow":"shadow","thick":"thick"}
    w_val = val_map.get(style, "single")

    for section in doc.sections:
        sectPr = section._sectPr
        for old in list(sectPr.findall(qn("w:pgBorders"))):
            sectPr.remove(old)
        pgBorders = OxmlElement("w:pgBorders")
        pgBorders.set(qn("w:offsetFrom"), "page")
        for edge in ("top", "left", "bottom", "right"):
            space = str(int(float(sizes.get(edge, 1.0)) * _CM_TO_TWENTIETHS))
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"),   w_val)
            el.set(qn("w:sz"),    "18")
            el.set(qn("w:space"), space)
            el.set(qn("w:color"), color)
            pgBorders.append(el)
        pgSz = sectPr.find(qn("w:pgSz"))
        if pgSz is not None:
            sectPr.insert(list(sectPr).index(pgSz), pgBorders)
        else:
            sectPr.append(pgBorders)


# ══════════════════════════════════════════════════════════════════════════════
#  TABLE CELL BORDERS
# ══════════════════════════════════════════════════════════════════════════════

def _apply_cell_borders(cell, style: str, color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in list(tcPr.findall(qn("w:tcBorders"))):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top","left","bottom","right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   style)
        tag.set(qn("w:sz"),    "4")
        tag.set(qn("w:color"), color.lstrip("#"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def apply_table_borders(table, style: str, color: str):
    seen = set()
    for row in table.rows:
        for cell in row.cells:
            if id(cell) not in seen:
                seen.add(id(cell))
                _apply_cell_borders(cell, style, color)


# ══════════════════════════════════════════════════════════════════════════════
#  LECTURE HOURS — right-aligned tab stop
# ══════════════════════════════════════════════════════════════════════════════

_LECTURE_HOURS_RE = re.compile(r"(\s*)(\(\d+\))\s*$")

def _add_right_tab_stop(para, position_inches: float = 6.0):
    """Add a right-aligned tab stop to the paragraph."""
    from docx.oxml import OxmlElement
    pPr  = para._p.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = OxmlElement("w:tabs")
        pPr.append(tabs)
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(int(position_inches * 1440)))   # twips
    tabs.append(tab)

def apply_lecture_hours_formatting(para, settings: dict):
    """
    Find (N) at end of paragraph. Insert tab before it and apply
    bold/italic based on settings. Also adds a right tab stop.
    """
    bold_h   = settings.get("bold_hours",         True)
    italic_h = settings.get("italic_hours",        False)
    right_h  = settings.get("right_align_hours",   True)

    full_text = para.text
    m = _LECTURE_HOURS_RE.search(full_text)
    if not m:
        return

    hours_text = m.group(2)          # e.g. "(9)"
    pre_text   = full_text[:m.start()]

    # Get font info from existing runs
    ref_font = None
    ref_size = None
    for run in para.runs:
        if run.text.strip():
            ref_font = run.font.name
            ref_size = run.font.size
            break

    # Rebuild: body text run + tab run + hours run
    p = para._p
    for r in list(p.findall(qn("w:r"))):
        p.remove(r)

    # Body text
    if pre_text.strip():
        body_run = para.add_run(pre_text.rstrip())
        if ref_font: body_run.font.name = ref_font
        if ref_size: body_run.font.size = ref_size

    # Tab (only if right-aligning)
    if right_h:
        tab_run = para.add_run("\t")
        if ref_font: tab_run.font.name = ref_font

    # Hours
    hours_run = para.add_run(hours_text)
    if ref_font: hours_run.font.name = ref_font
    if ref_size: hours_run.font.size = ref_size
    hours_run.bold   = bold_h
    hours_run.italic = italic_h

    # Add tab stop at right margin
    if right_h:
        _add_right_tab_stop(para, position_inches=6.0)


# ══════════════════════════════════════════════════════════════════════════════
#  RUN FONT HELPER
# ══════════════════════════════════════════════════════════════════════════════

def _set_run_font(run, font_name: str, font_pt, bold=None, italic=None):
    run.font.name = font_name
    run.font.size = font_pt
    if bold   is not None: run.bold   = bold
    if italic is not None: run.italic = italic


# ══════════════════════════════════════════════════════════════════════════════
#  PARAGRAPH CLASSIFIERS
# ══════════════════════════════════════════════════════════════════════════════

_UNIT_HEADING_RE = re.compile(r"^[A-Z][A-Z\s,/]+:\s")
_REF_KEYS        = ["TEXT BOOKS", "TEXTBOOKS", "REFERENCES"]
_END_KEYS        = ["COURSE OUTCOMES", "MAPPING COURSE", "ARTICULATION", "TOTAL L"]

def _is_unit_heading(para)  -> bool: return bool(_UNIT_HEADING_RE.match(para.text.strip()))
def _is_ref_start(upper)    -> bool: return any(k in upper for k in _REF_KEYS)
def _is_ref_end(upper)      -> bool: return any(k in upper for k in _END_KEYS)
def _has_lecture_hours(para)-> bool: return bool(_LECTURE_HOURS_RE.search(para.text))


# ══════════════════════════════════════════════════════════════════════════════
#  REFERENCE EXTRACTOR
# ══════════════════════════════════════════════════════════════════════════════

def extract_references(docx_path: str) -> list:
    doc = Document(docx_path)
    refs, in_ref = [], False
    for para in doc.paragraphs:
        text  = para.text.strip()
        upper = text.upper()
        if _is_ref_start(upper):  in_ref = True;  continue
        if in_ref:
            if _is_ref_end(upper): in_ref = False; continue
            if text and len(text) > 10: refs.append(text)
    return refs


# ══════════════════════════════════════════════════════════════════════════════
#  GROQ RUN-LEVEL CLASSIFICATION (topic headings inside body paragraphs)
# ══════════════════════════════════════════════════════════════════════════════

def apply_groq_run_classification(para, classified_runs: list, settings: dict):
    uh_font  = settings.get("unit_heading_font",  "Times New Roman")
    uh_pt    = Pt(settings.get("unit_heading_size", 12))
    th_font  = settings.get("topic_heading_font", settings.get("body_font", "Times New Roman"))
    th_pt    = Pt(settings.get("topic_heading_size", settings.get("body_size", 11)))
    body_font= settings.get("body_font",  "Times New Roman")
    body_pt  = Pt(settings.get("body_size", 11))
    ref_font = settings.get("reference_font", "Times New Roman")
    ref_pt   = Pt(settings.get("reference_size", 11))
    tbl_font = settings.get("table_header_font", "Arial")
    tbl_pt   = Pt(settings.get("table_header_size", 10))
    bold_top = settings.get("bold_topic_headings", True)

    TYPE_MAP = {
        "unit_heading":  (uh_font,   uh_pt,   True,      None),
        "topic_heading": (th_font,   th_pt,   bold_top,  None),
        "section_label": (uh_font,   uh_pt,   True,      None),
        "body":          (body_font, body_pt, None,      None),
        "lecture_hours": (body_font, body_pt, True,      None),
        "reference":     (ref_font,  ref_pt,  None,      None),
        "table_header":  (tbl_font,  tbl_pt,  True,      None),
        "table_body":    (tbl_font,  tbl_pt,  None,      None),
    }

    p = para._p
    for r in list(p.findall(qn("w:r"))):
        p.remove(r)

    for cr in classified_runs:
        rtype = cr.get("type", "body")
        fn, fp, bd, it = TYPE_MAP.get(rtype, (body_font, body_pt, None, None))
        run = para.add_run(cr.get("text", ""))
        _set_run_font(run, fn, fp, bd, it)
        if cr.get("italic"): run.italic = True


# ══════════════════════════════════════════════════════════════════════════════
#  REFERENCE PART FORMATTING
# ══════════════════════════════════════════════════════════════════════════════

def apply_reference_parts(para, ref_text: str, settings: dict,
                          groq_api_key: str, ref_font: str, ref_pt):
    """Format one reference line with per-part bold/italic via Groq."""
    from groq_parser import format_reference_parts

    ref_format = {k: settings.get(k, False)
                  for k in ("author_italic","author_bold","title_italic","title_bold",
                             "publisher_italic","publisher_bold","edition_italic",
                             "edition_bold","year_italic","year_bold")}

    parts = format_reference_parts(ref_text, groq_api_key, ref_format)

    p = para._p
    for r in list(p.findall(qn("w:r"))):
        p.remove(r)
    for part in parts:
        run = para.add_run(part["text"])
        _set_run_font(run, ref_font, ref_pt)
        run.bold   = part.get("bold",   False)
        run.italic = part.get("italic", False)


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN FORMATTER
# ══════════════════════════════════════════════════════════════════════════════

def format_document(
    input_path:       str,
    output_path:      str,
    settings:         dict,
    cleaned_refs:     list = None,
    groq_api_key:     str  = None,
    use_run_classify: bool = False,
):
    from groq_parser import classify_paragraph_runs

    doc = Document(input_path)

    # ── Page ──────────────────────────────────────────────────────────────────
    apply_page_size(doc, settings.get("page_size", "A4"))
    apply_page_border(
        doc,
        settings.get("page_border_style", "none"),
        settings.get("page_border_color", "000000"),
        {e: settings.get(f"page_border_{e}", 1.0)
         for e in ("top","bottom","left","right")},
    )

    # ── Font objects ──────────────────────────────────────────────────────────
    uh_font  = settings.get("unit_heading_font",  "Times New Roman")
    uh_pt    = Pt(settings.get("unit_heading_size", 12))
    body_font= settings.get("body_font",  "Times New Roman")
    body_pt  = Pt(settings.get("body_size", 11))
    th_font  = settings.get("table_header_font", "Arial")
    th_pt    = Pt(settings.get("table_header_size", 10))
    tb_font  = settings.get("table_body_font",   "Arial")
    tb_pt    = Pt(settings.get("table_body_size", 10))
    ref_font = settings.get("reference_font", "Times New Roman")
    ref_pt   = Pt(settings.get("reference_size", 11))

    # ── Alignment ─────────────────────────────────────────────────────────────
    _align = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "left":    WD_ALIGN_PARAGRAPH.LEFT,
        "center":  WD_ALIGN_PARAGRAPH.CENTER,
        "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    }
    body_align = _align.get(settings.get("alignment","justify"),
                             WD_ALIGN_PARAGRAPH.JUSTIFY)

    in_ref = False
    ref_idx = 0

    # ── Paragraphs ────────────────────────────────────────────────────────────
    for para in doc.paragraphs:
        text  = para.text.strip()
        upper = text.upper()

        if _is_ref_start(upper):  in_ref = True
        elif _is_ref_end(upper):  in_ref = False
        if not text: continue

        has_hours = _has_lecture_hours(para)

        # ── Groq run-level classification ─────────────────────────────────
        if use_run_classify and groq_api_key and not in_ref:
            runs_meta = [{"text": r.text, "bold": bool(r.bold), "italic": bool(r.italic)}
                         for r in para.runs if r.text]
            if runs_meta:
                classified = classify_paragraph_runs(runs_meta, groq_api_key)
                apply_groq_run_classification(para, classified, settings)
                if has_hours:
                    apply_lecture_hours_formatting(para, settings)
                continue

        # ── Unit heading ──────────────────────────────────────────────────
        if _is_unit_heading(para):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                if run.bold or run.text.strip().isupper():
                    _set_run_font(run, uh_font, uh_pt,
                                  bold=settings.get("bold_headings", True))
                else:
                    _set_run_font(run, body_font, body_pt)
            if has_hours:
                apply_lecture_hours_formatting(para, settings)

        # ── Reference entry ───────────────────────────────────────────────
        elif in_ref:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            is_hdr = any(k in upper for k in _REF_KEYS)
            if is_hdr:
                for run in para.runs:
                    _set_run_font(run, uh_font, uh_pt, bold=True)
                continue

            # Get the (possibly cleaned) text
            ref_text = text
            if cleaned_refs and ref_idx < len(cleaned_refs):
                ref_text = cleaned_refs[ref_idx]
                ref_idx += 1

            if groq_api_key:
                apply_reference_parts(para, ref_text, settings,
                                      groq_api_key, ref_font, ref_pt)
            else:
                # Fallback: whole line, italic quoted titles
                p = para._p
                for r in list(p.findall(qn("w:r"))): p.remove(r)
                run = para.add_run(ref_text)
                _set_run_font(run, ref_font, ref_pt)

        # ── Body ──────────────────────────────────────────────────────────
        else:
            para.alignment = body_align
            for run in para.runs:
                _set_run_font(run, body_font, body_pt)
            if has_hours:
                apply_lecture_hours_formatting(para, settings)

    # ── Tables ────────────────────────────────────────────────────────────────
    cbs = settings.get("cell_border_style", "single")
    cbc = settings.get("cell_border_color", "000000")
    for table in doc.tables:
        if cbs != "none":
            apply_table_borders(table, cbs, cbc)
        for row_idx, row in enumerate(table.rows):
            is_hdr = (row_idx == 0)
            fn = th_font if is_hdr else tb_font
            fp = th_pt   if is_hdr else tb_pt
            seen = set()
            for cell in row.cells:
                if id(cell) in seen: continue
                seen.add(id(cell))
                for cp in cell.paragraphs:
                    for run in cp.runs:
                        _set_run_font(run, fn, fp,
                                      bold=True if is_hdr else run.bold)

    doc.save(output_path)
    return True
